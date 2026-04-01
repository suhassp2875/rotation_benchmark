import docx
from docx.shared import Inches
import os

def add_heading_and_paragraph(doc, heading, text, level=1):
    doc.add_heading(heading, level=level)
    if text:
        doc.add_paragraph(text)

def main():
    doc = docx.Document()
    
    doc.add_heading('RotationBench CoRL Paper Scaffold', 0)
    
    # 1. Core Narrative
    narrative = (
        "This is an empirical robot-learning paper focused on reliability. "
        "The central thesis is that standard regression metrics (like median geodesic error) "
        "mask catastrophic, rare orientation failures—'tail risk'—that critically degrade downstream robotic manipulation. "
        "By formalizing worst-case metrics (p99 error and failure rates >30 degrees) and evaluating "
        "across a 25M-sample synthetic benchmark and controlled simulated robotic tasks, we demonstrate that a representation's "
        "reliability at the antipodal boundary (pi) is strongly predictive of its practical safety in control pipelines in this setting. "
        "We establish that continuous methods (like 6D and SVD) form a reliable top tier, while minimal variants require "
        "specific architectural hardening to suppress high-variance tail failures."
    )
    add_heading_and_paragraph(doc, '1. Core Narrative', narrative)
    
    # 2. Title Options
    add_heading_and_paragraph(doc, '2. Title Options', None)
    doc.add_paragraph("• Evaluating Rotation Parameterizations for Reliable Robot Learning: A Focus on Worst-Case Failures", style='List Bullet')
    doc.add_paragraph("• Beyond Median Error: Characterizing the Tail-Risk of Rotation Representations in Robot Learning", style='List Bullet')
    doc.add_paragraph("• Reliable Rotation Regression for Robotics: An Empirical Study of Representation Tail-Risk", style='List Bullet')
    
    # 3. Abstract
    abstract = (
        "While modern robot learning pipelines heavily rely on 3D rotation parameterizations, these representations are typically evaluated on average-case accuracy. "
        "In an embodied setting, however, rare but catastrophic orientation errors can lead to task-critical failures. "
        "In this work, we present an empirical investigation into the worst-case reliability ('tail-risk') of neural rotation representations. "
        "Through a large-scale (25M-sample) synthetic benchmark, we show that while many standard representations achieve low median geodesic errors, "
        "they exhibit drastically different tail behaviors, with severe failures concentrating near the antipodal (pi) boundary. "
        "We formalize p99 and threshold bounding (>30 degrees) as primary reliability metrics. "
        "To ground these findings in physical task success, we conduct a controlled error-injection evaluation in simulated robotic manipulation (Robosuite). "
        "Our results demonstrate a sharp 'cliff zone' where task success degrades rapidly as orientation errors enter the upper tail. "
        "We find that continuous mappings (e.g., 6D and SVD) form a robust top tier for reliable regression, while minimal Lie-family variants require "
        "specific architectural hardening to mitigate boundary ambiguity. "
        "Finally, through live model-in-the-loop evaluations, we provide evidence that lower tail-risk is associated with higher downstream reliability in controlled manipulation tasks."
    )
    add_heading_and_paragraph(doc, '3. Abstract Draft', abstract)
    
    # 4. Full Paper Skeleton
    add_heading_and_paragraph(doc, '4. Full Paper Skeleton', None)
    
    skel = [
        ("1. Introduction", "State the problem: robots fail when rotations are catastrophically wrong. Claim: Tail-risk metrics better predict downstream reliability. Avoid: Claiming we solved rotation learning. [SOFTEN CLAIM] on robotics—ensure we specify simulated downstream tasks."),
        ("2. Background & Problem Formulation", "Define the supervised mapping, geodesic distance, q_0.99, and Fail_30. Observation: Near rotation angle pi, minimal Euclidean parameterizations admit multiple target encodings, creating unstable boundary supervision. Avoid: Over-theorizing."),
        ("3. The RotationBench Evaluation Framework", "Describe the 25M synthetic setup, shared backbone, and evaluated representations. Claim: This controlled setup isolates representation-induced failure modes. Avoid: Claiming the simple MLP backbone is deployment-ready for arbitrary architectures."),
        ("4. Empirical Reliability Analysis (Synthetic)", "Present the core synthetic findings and the failure taxonomy. FullFix acts as the headline minimal variant. Claim: Continuous methods (SVD, 6D) form the reliable top tier. Failures concentrate near the hard-angle region. Avoid: Claiming SVD is universally superior to 6D."),
        ("5. Robotics-Grounded Downstream Validation", "Connect abstract metrics to physical task success. Includes 5.1 Controlled Error-Injection, 5.2 The Tail-Risk Cliff Zone, 5.3 Multi-task Evaluation, and 5.4 Live Model-in-the-Loop (Main Paper). Claim: Lower tail error is associated with higher downstream reliability. Avoid: Saying this guarantees real-world safety."),
        ("6. Limitations & Future Work", "Honest boundary-setting. Highlight that BOP / YCB-V transfer was not successful due to the simple order-dependent MLP backbone. State that real-world deployment requires permutation-invariant architectures like PointNet. Explicitly keep this as a limitation."),
        ("7. Conclusion", "Crisp summary.")
    ]
    for title, text in skel:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
        
    # 5. Figure Plan
    add_heading_and_paragraph(doc, '5. Figure Plan', None)
    
    fig_plan = [
        ("Figure 1 (Main): Main Synthetic Benchmark Summary", "Visual comparison of median vs tail performance across the core headline methods to establish the benchmark's central finding right away.", ""),
        ("Figure 2 (Main): Failure Taxonomy & Tail Risk Concept", "Show error distributions split into Stable (<5), Moderate (5-30), and Severe (>30).", "results/figures/failure_taxonomy.png"),
        ("Figure 3 (Main): The Error Sweep 'Cliff Zone'", "Line plot showing task success drops sharply as injected error scales.", "results/figures/error_sweep.png"), # May not exist exactly, but we add placeholder if needed
        ("Figure 4 (Main): Multi-Task Simulation Snapshots", "Visual context for PickPlace, Stack, NutAssembly.", "results/figures/sim_panel_pickplace.png"),
        ("Figure 5 (Main/Appendix): Antipodal Boundary Heatmap", "Prove that failures map to pi boundary.", "results/figures/boundary_stress.png")
    ]
    
    for title, desc, img_path in fig_plan:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)
        if img_path and os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(5.0))
            except Exception as e:
                doc.add_paragraph(f"[Image {img_path} not available to embed: {e}]")
        elif img_path:
            doc.add_paragraph(f"[Image planned: {os.path.basename(img_path)}]")

    # 6. Table Plan
    add_heading_and_paragraph(doc, '6. Table Plan', None)
    doc.add_heading("Table 1 (Main): Synthetic Benchmark Results", level=2)
    doc.add_paragraph("Rows: SVD (1.95, 4.27, 0%), 6D (1.81, 3.77, 0%), Quat (3.98, 57.39, 2.3%), Atlas [Oracle] (3.33, 28.76, 0.9%), Lie_FullFix (3.65, 57.27, 3.5%).")
    
    doc.add_heading("Table 2 (Main): Lie-Family Ablation", level=2)
    doc.add_paragraph("Rows: Raw Lie, Hardened, FullFix. Shows how architectural hardening alters the central/tail trade-off.")
    
    doc.add_heading("Table 3 (Main): Robosuite Multi-Task Summary", level=2)
    doc.add_paragraph("Rows: Baseline, SVD, 6D, Atlas, Quat, Lie_FullFix. Evaluates PickPlace, Stack, and NutAssembly. NutAssembly is a hard-ceiling diagnostic.")
    
    doc.add_heading("Table 4 (Main): Live Model-in-the-loop Results", level=2)
    doc.add_paragraph("Promoted to main paper. PickPlace live evaluation with SVD, 6D, FullFix, Raw Lie proving the dynamic case matches static error injection insights.")
    
    # 7. Safe Claim Language
    add_heading_and_paragraph(doc, '7. Reviewer-Safe Claim Language', None)
    doc.add_paragraph("• 'We provide controlled empirical evidence that...'")
    doc.add_paragraph("• 'Our synthetic benchmark isolates representation-induced failure modes...'")
    doc.add_paragraph("• 'Continuous methods (SVD, 6D) form a reliable top tier for bounding tail-risk...'")
    doc.add_paragraph("• 'Across achievable evaluated simulated tasks, lower tail error is associated with higher downstream robustness...'")
    doc.add_paragraph("• 'We identify the hard-angle boundary as a primary source of representation instability...'")
    
    # 8. What Not To Oversell
    add_heading_and_paragraph(doc, '8. What Not to Oversell', None)
    doc.add_paragraph("• Do not say 'SVD is universally superior.' -> Say 'SVD and 6D provide top-tier reliability.'")
    doc.add_paragraph("• Do not say 'Guarantees safe real-world robot deployment.' -> Say 'Bounding tail-risk is predictive of improved task success in continuous control simulations.'")
    doc.add_paragraph("• Do not hide the Atlas oracle caveat, the MLP external transfer failure (BOP), or the NutAssembly hard-ceiling.")
    
    doc.save("RotationBench_CoRL_Scaffold.docx")
    print("Document successfully created: RotationBench_CoRL_Scaffold.docx")

if __name__ == '__main__':
    main()

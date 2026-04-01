"""
generate_paper_docx.py
Generates a full research paper in .docx format including all figures and tables.
"""
import os, json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

# Paths
REPO_ROOT = Path("c:/Users/suhas/Downloads/rotation-benchmark/rotation-benchmark")
FIG_DIR   = REPO_ROOT / "results/paper_figures"
GOLD_JSON = REPO_ROOT / "results/gold_standard.json"
LIVE_JSON = REPO_ROOT / "results/live_eval_results.json"
OUT_FILE  = REPO_ROOT / "Evaluation_of_Rotation_Parameterizations_for_Reliable_Robot_Learning.docx"

def add_centered_image(doc, img_path, caption, width=Inches(5.0)):
    if not img_path.exists():
        print(f"  [WARN] Image not found: {img_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=width)
    cap = doc.add_paragraph(f"Figure: {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = 'Caption' if 'Caption' in doc.styles else None

def main():
    print("Generating Paper Draft...")
    doc = Document()
    
    # --- Title ---
    title = doc.add_heading('Evaluating Rotation Parameterizations for Reliable Robot Learning: A Focus on Worst-Case Failures', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "Supervised learning of 3-D rotations is a fundamental component in modern robotics, serving roles in object pose estimation, "
        "end-effector control, and visual odometry. While prevailing benchmarks emphasize average-case performance metrics, "
        "these summaries can obscure rare but operationally catastrophic failures. In this work, we present a large-scale evaluation "
        "centered on worst-case reliability, showing that lower p99 error is associated with improved downstream manipulation success. "
        "By comparing continuous higher-dimensional mappings, minimal Lie-algebraic vectors, and multi-hypothesis models, "
        "we find that top-tier continuous methods (SVD/6D) dominate worst-case robustness. Our findings suggest that rotation "
        "representations should be judged by their tail-risk reliability in safety-critical robotics settings."
    )

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Supervised learning of 3D rotations is a fundamental component in modern robotics, with applications in object pose estimation, "
        "end-effector control, and visual odometry. While most benchmarks emphasize average-case metrics such as median geodesic error, "
        "these summaries can obscure rare but operationally catastrophic failures, particularly near antipodal ambiguities."
    )
    doc.add_paragraph(
        "In this work, we present a large-scale evaluation centered on worst-case reliability, showing that lower p99 error is associated "
        "with improved downstream manipulation success in controlled robotics settings."
    )

    # --- Related Work ---
    doc.add_heading('2. Related Work', level=1)
    doc.add_paragraph(
        "The choice of rotation parameterization has long been recognized as a key factor in the stability and efficiency of "
        "learning-based pose prediction. While continuous higher-dimensional mappings such as 6D offer favorable optimization properties, "
        "minimal representations remain attractive because of their compactness, motivating continued work on ambiguity-aware losses and geometry-aware supervision."
    )

    # --- Methods ---
    doc.add_heading('3. Methods', level=1)
    doc.add_paragraph(
        "We compare continuous representations, minimal Lie-algebraic mappings, and chart-based multi-hypothesis models under a matched training "
        "and evaluation protocol. To probe whether loss design can mitigate antipodal instability in minimal mappings, we evaluate a FullFix variant "
        "that combines ambiguity-aware supervision with architectural hardening."
    )
    
    # --- Experimental Setup ---
    doc.add_heading('4. Experimental Setup', level=1)
    doc.add_paragraph(
        "Our synthetic benchmark consists of 25 million sample pairs with broad SO(3) coverage, emphasizing the hard-angle regime near the antipodal boundary. "
        "We prioritize worst-case metrics, specifically the p99 geodesic error and the failure rate above 30 degrees."
    )

    # --- Results ---
    doc.add_heading('5. Main Benchmark Results', level=1)
    doc.add_paragraph(
        "Across the synthetic benchmark, low-tail continuous methods form the clear high-reliability tier. Although median errors "
        "are relatively close across methods, tail behavior differs sharply."
    )
    
    # Table 1: Main Benchmark
    if GOLD_JSON.exists():
        with open(GOLD_JSON, "r") as f: gold = json.load(f)
        headline = gold.get("headline", {})
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Representation'
        hdr_cells[1].text = 'Median Error (°)'
        hdr_cells[2].text = 'p99 Error (°)'
        hdr_cells[3].text = 'Fail Rate (>30°)'
        
        for name, m in headline.items():
            row_cells = table.add_row().cells
            row_cells[0].text = name.upper()
            row_cells[1].text = f"{m['median']:.2f}"
            row_cells[2].text = f"{m['p99']:.2f}"
            row_cells[3].text = f"{m['fail30']:.1f}%"
        doc.add_paragraph("Table 1. Main synthetic benchmark results (N=5000).").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_centered_image(doc, FIG_DIR / "fig1_p99_vs_task_success.png", "p99 predicts robot success.")
    add_centered_image(doc, FIG_DIR / "fig7_error_sweep_curve.png", "Monotonic degradation and the 'Cliff Zone'.")

    # --- Failure Anatomy ---
    doc.add_heading('6. Failure Anatomy', level=1)
    doc.add_paragraph(
        "Error stratified by ground-truth rotation magnitude reveals that failures are concentrated near the antipodal boundary. "
        "Ablation results suggest that specialized objectives can reduce tail failures, though they do not fully remove representational fragility."
    )
    add_centered_image(doc, FIG_DIR / "fig6_failure_taxonomy.png", "Failure taxonomy across error zones.")

    # --- Robotics Validation ---
    doc.add_heading('7. Downstream Robotics Validation', level=1)
    doc.add_paragraph(
        "To study the operational consequences of error patterns, we evaluate performance in robosuite. Percentile-based sweeps "
        "reveal a 'cliff zone' where task success degrades sharply once errors enter the upper tail of a method's distribution."
    )
    
    # Table: Live Eval
    if LIVE_JSON.exists():
        with open(LIVE_JSON, "r") as f: live = json.load(f)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Model'
        hdr_cells[1].text = 'Live Success Rate'
        for name, m in live.items():
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = f"{m['success_rate']*100:.1f}%"
        doc.add_paragraph("Table 2. Live model-in-the-loop success rates.").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_centered_image(doc, FIG_DIR / "fig5a_sim_panel_pickplace.png", "Qualitative simulation outcomes.")

    # --- Conclusion ---
    doc.add_heading('8. Conclusion', level=1)
    doc.add_paragraph(
        "Our findings provide strong evidence that p99 error is a useful predictor of downstream task reliability. "
        "For practitioners, these findings support using low-tail continuous representations as strong default choices."
    )

    doc.save(OUT_FILE)
    print(f"Successfully saved paper to {OUT_FILE}")

if __name__ == "__main__":
    main()
